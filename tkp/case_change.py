import pymorphy2
from docx import Document
import nltk

# Инициализация
morph = pymorphy2.MorphAnalyzer()
oughtList = [item.word for item in morph.parse("должны")[0].lexeme]


def convert(line):
    # print("Исходная строка")
    # print(line.replace('\r', '').replace('\n', ''))
    n = -1
    # Ищем слово должен
    for item in oughtList:
        try:
            n = line.lower().split().index(item)
            break
        except ValueError:
            pass
    # Нашли слово
    if n >= 0:
        newLine = line.split()
        src = morph.parse(newLine[n])[0]
        next = morph.parse(newLine[n + 1])[0]
        # если совершенный вид удаляем "должен"
        if next.tag.aspect == 'perf':
            newLine[n + 1] = next.inflect({'3per'}).word
            newLine.pop(n)
        else:
            will = morph.parse("будет")[0]
            replace = will.inflect({src.tag.number}) if src.tag.number else will
            newLine[n] = replace.word
            # если следующее слово - форма слова быть
            if newLine[n + 1] in [item.word for item in will.lexeme]:
                newLine.pop(n + 1)
        print(line)
        print(" ".join(newLine))
        return " ".join(newLine)
    else:
        return line

    # print("Результат")
    # print(*newLine)
    # print()


# with open("D:\\temp\\text.txt", "r") as f:
#    for line in f:
#        print("Исходная строка")
#        print(line.replace('\r', '').replace('\n', ''))
#        print("Результат")
#        print(convert(line))
#        print()



def convert_file(infile, outfile):
    document = Document(infile)

    for item in document.paragraphs:
        sentences = nltk.tokenize.sent_tokenize(item.text)
        # sentences = split_into_sentences(item.text)
        # sentences = item.text.split('.')
        # print(sentences)
        sentences = map(convert, sentences)
        item.text = " ".join(sentences)
        # print(item.text)
        # print(sentences)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for item in cell.paragraphs:
                    sentences = nltk.tokenize.sent_tokenize(item.text)
                    # sentences = split_into_sentences(item.text)
                    # sentences = item.text.split('.')
                    # print(sentences)
                    sentences = map(convert, sentences)
                    item.text = " ".join(sentences)

    document.save(outfile)
